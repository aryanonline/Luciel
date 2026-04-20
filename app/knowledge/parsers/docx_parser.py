"""Microsoft Word .docx parser (Step 25b, File 5)."""
from __future__ import annotations

from io import BytesIO

import docx as docx_lib
from docx.opc.exceptions import PackageNotFoundError

from app.knowledge.parsers.base import ParsedDocument, Parser, ParserError


class DocxParser(Parser):
    source_type = "docx"

    def parse(self, file_bytes: bytes, filename: str | None = None) -> ParsedDocument:
        try:
            document = docx_lib.Document(BytesIO(file_bytes))
        except PackageNotFoundError as exc:
            raise ParserError(f"Not a valid .docx file: {exc}") from exc

        paragraphs: list[str] = []
        for p in document.paragraphs:
            txt = p.text.strip()
            if txt:
                paragraphs.append(txt)

        # Include tables — one cell per line, row breaks as single newlines,
        # table breaks as double newlines so the paragraph chunker respects them.
        table_blocks: list[str] = []
        for table in document.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                table_blocks.append("\n".join(rows))

        all_blocks = paragraphs + table_blocks
        if not all_blocks:
            raise ParserError(".docx document produced no extractable text")
        text = "\n\n".join(all_blocks)
        return ParsedDocument(
            text=text,
            metadata={
                "paragraph_count": len(paragraphs),
                "table_count": len(document.tables),
            },
        )